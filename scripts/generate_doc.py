#!/usr/bin/env python3
"""
政府公文生成器 - 按《党政机关公文格式》(GB/T 9704-2012) 标准生成 Word 文档。

用法:
  python generate_doc.py --title "标题" --recipient "主送机关" \
    --body-file content.json --sender "发文机关" --date "2026年5月22日" \
    --output output.docx
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)


# ========== 格式常量（EMU 单位） ==========
# 页面设置
A4_WIDTH = 11906   # 210mm in TWIPS (for python-docx section)
A4_HEIGHT = 16838  # 297mm in TWIPS

# 字号 (EMU)
FONT_SIZE_TITLE = Pt(22)     # 二号
FONT_SIZE_BODY = Pt(16)      # 三号

# 行距固定值 (EMU)
LINE_SPACING_TITLE = Pt(38)  # 标题行
LINE_SPACING_BODY = Pt(28)   # 正文行

# 缩进
INDENT_2CHAR = Pt(32)        # 首行缩进2字符 ≈ 406400 EMU
INDENT_SENDER = Pt(160)      # 落款署名缩进
INDENT_DATE = Pt(224)        # 成文日期缩进

# 附件缩进
ATTACH_LEFT_INDENT = Pt(104)      # 左缩进
ATTACH_HANGING_INDENT = Pt(-72)   # 悬挂缩进


def set_line_spacing(paragraph, spacing, rule_type="exactly"):
    """设置段落行距为固定值"""
    pf = paragraph.paragraph_format
    pf.line_spacing = spacing
    from docx.enum.text import WD_LINE_SPACING
    if rule_type == "exactly":
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def add_paragraph(doc, text, font_name_cn, font_name_en, font_size, alignment=None,
                  first_line_indent=None, line_spacing=None, bold=False,
                  left_indent=None, right_indent=None):
    """添加格式化的段落"""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    if alignment is not None:
        para.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if right_indent is not None:
        pf.right_indent = right_indent
    set_line_spacing(para, line_spacing or LINE_SPACING_BODY)

    run = para.add_run(text)
    run.font.size = font_size
    run.font.bold = bold

    # 设置中文字体和西文字体
    run.font.name = font_name_cn
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    rPr.insert(0, rFonts)

    return para


def add_body_paragraph(doc, text, first_line_indent=True):
    """添加正文段落（仿宋，16pt，首行缩进2字符）"""
    indent = INDENT_2CHAR if first_line_indent else None
    return add_paragraph(
        doc, text,
        font_name_cn="仿宋",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=indent,
        line_spacing=LINE_SPACING_BODY
    )


def add_heading1(doc, text):
    """添加一级标题：黑体，16pt，首行缩进2字符"""
    return add_paragraph(
        doc, text,
        font_name_cn="黑体",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=INDENT_2CHAR,
        line_spacing=LINE_SPACING_BODY,
        bold=True
    )


def add_heading2(doc, text):
    """添加二级标题：楷体，16pt，首行缩进2字符"""
    return add_paragraph(
        doc, text,
        font_name_cn="楷体",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=INDENT_2CHAR,
        line_spacing=LINE_SPACING_BODY,
        bold=True
    )


def add_heading3(doc, text):
    """添加三级标题：仿宋，16pt，首行缩进2字符"""
    return add_paragraph(
        doc, text,
        font_name_cn="仿宋",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=INDENT_2CHAR,
        line_spacing=LINE_SPACING_BODY
    )


def add_heading4(doc, text):
    """添加四级标题：仿宋，16pt，首行缩进2字符"""
    return add_paragraph(
        doc, text,
        font_name_cn="仿宋",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=INDENT_2CHAR,
        line_spacing=LINE_SPACING_BODY
    )


def add_empty_line(doc):
    """添加空行"""
    return add_body_paragraph(doc, "", first_line_indent=False)


def generate_document(title, main_recipient, body_items, sender, date_str, attachments=None):
    """生成公文文档"""
    doc = Document()

    # ====== 页面设置 ======
    section = doc.sections[0]
    section.page_width = Cm(21)     # A4 210mm
    section.page_height = Cm(29.7)  # A4 297mm
    section.top_margin = Cm(3.7)    # 37mm
    section.bottom_margin = Cm(3.5) # 35mm
    section.left_margin = Cm(2.8)   # 28mm
    section.right_margin = Cm(2.6)  # 26mm

    # ====== 1. 标题 ======
    add_paragraph(
        doc, title,
        font_name_cn="方正小标宋_GBK",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_TITLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=LINE_SPACING_TITLE
    )

    # ====== 2. 空一行（预留红头间距） ======
    add_empty_line(doc)

    # ====== 3. 主送机关 ======
    add_paragraph(
        doc, main_recipient,
        font_name_cn="仿宋",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        line_spacing=LINE_SPACING_BODY
    )

    # ====== 4. 正文内容 ======
    for item in body_items:
        item_type = item.get("type", "text")
        content = item.get("content", "")

        if item_type == "text":
            for para in content.split("\n"):
                para = para.strip()
                if para:
                    add_body_paragraph(doc, para)
                else:
                    add_empty_line(doc)
        elif item_type == "heading1":
            add_heading1(doc, content)
        elif item_type == "heading2":
            add_heading2(doc, content)
        elif item_type == "heading3":
            add_heading3(doc, content)
        elif item_type == "heading4":
            add_heading4(doc, content)
        elif item_type == "empty":
            add_empty_line(doc)

    # ====== 5. 附件说明 ======
    if attachments and len(attachments) > 0:
        add_empty_line(doc)

        # 第一行附件标题
        first_attach_text = f"附件：1．{attachments[0]}"
        add_paragraph(
            doc, first_attach_text,
            font_name_cn="仿宋",
            font_name_en="Times New Roman",
            font_size=FONT_SIZE_BODY,
            left_indent=ATTACH_LEFT_INDENT,
            first_line_indent=ATTACH_HANGING_INDENT,
            line_spacing=LINE_SPACING_BODY
        )

        # 后续附件
        for i, attach_name in enumerate(attachments[1:], start=2):
            attach_text = f"{i}．{attach_name}"
            add_paragraph(
                doc, attach_text,
                font_name_cn="仿宋",
                font_name_en="Times New Roman",
                font_size=FONT_SIZE_BODY,
                left_indent=ATTACH_LEFT_INDENT,
                first_line_indent=Pt(-52),  # 略小于第一行的悬挂
                line_spacing=LINE_SPACING_BODY
            )

    # ====== 6. 空行 ======
    add_empty_line(doc)

    # ====== 7. 发文机关署名 ======
    add_paragraph(
        doc, sender,
        font_name_cn="仿宋",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=INDENT_SENDER,
        line_spacing=LINE_SPACING_BODY
    )

    # ====== 8. 成文日期 ======
    add_paragraph(
        doc, date_str,
        font_name_cn="仿宋",
        font_name_en="Times New Roman",
        font_size=FONT_SIZE_BODY,
        first_line_indent=INDENT_DATE,
        line_spacing=LINE_SPACING_BODY
    )

    return doc


def main():
    parser = argparse.ArgumentParser(description="政府公文生成器")
    parser.add_argument("--title", required=True, help="公文标题")
    parser.add_argument("--recipient", required=True, help="主送机关名称")
    parser.add_argument("--body-file", required=True, help="正文内容 JSON 文件路径")
    parser.add_argument("--sender", required=True, help="发文机关署名")
    parser.add_argument("--date", required=True, help="成文日期")
    parser.add_argument("--attachments", default=None, help="附件名称列表，逗号分隔")
    parser.add_argument("--output", default="output.docx", help="输出文件路径")

    args = parser.parse_args()

    # 读取正文 JSON
    with open(args.body_file, "r", encoding="utf-8") as f:
        body_items = json.load(f)

    # 解析附件
    attachments = None
    if args.attachments:
        attachments = [a.strip() for a in args.attachments.split(",") if a.strip()]

    # 生成文档
    doc = generate_document(
        title=args.title,
        main_recipient=args.recipient,
        body_items=body_items,
        sender=args.sender,
        date_str=args.date,
        attachments=attachments
    )

    # 保存
    output_path = Path(args.output)
    doc.save(str(output_path))
    print(f"公文已生成: {output_path.absolute()}")


if __name__ == "__main__":
    main()
